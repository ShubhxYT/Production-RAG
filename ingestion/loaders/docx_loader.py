"""DOCX document loader using python-docx."""

import logging
import zipfile
from pathlib import Path

from docx import Document as DocxFile
from docx.table import Table
from docx.text.paragraph import Paragraph

from ingestion.models import Document

logger = logging.getLogger(__name__)


class DocxLoader:
    """Load and convert DOCX files to markdown."""

    def load(self, file_path: Path, output_dir: Path) -> Document:
        """Load a DOCX file and convert it to a Document with markdown content.

        Args:
            file_path: Path to the DOCX file.
            output_dir: Directory to save extracted images.

        Returns:
            Document with raw_content populated as markdown text.

        Raises:
            FileNotFoundError: If the DOCX file does not exist.
        """
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        doc_output_dir = output_dir / file_path.stem
        doc_output_dir.mkdir(parents=True, exist_ok=True)

        try:
            docx = DocxFile(str(file_path))
        except Exception:
            logger.exception("Failed to open DOCX: %s", file_path)
            raise

        # Build element map for ordered body iteration
        element_map: dict = {}
        for para in docx.paragraphs:
            element_map[para._element] = para
        for table in docx.tables:
            element_map[table._element] = table

        md_parts: list[str] = []
        title: str | None = None

        for child in docx.element.body:
            obj = element_map.get(child)
            if obj is None:
                continue

            if isinstance(obj, Paragraph):
                text = obj.text.strip()
                if not text:
                    continue

                style_name = obj.style.name if obj.style else "Normal"

                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    md_parts.append(f"{'#' * level} {text}")
                    if title is None:
                        title = text
                elif style_name == "Title":
                    md_parts.append(f"# {text}")
                    if title is None:
                        title = text
                elif style_name.startswith("List"):
                    md_parts.append(f"- {text}")
                else:
                    md_parts.append(text)

            elif isinstance(obj, Table):
                md_parts.append(self._table_to_markdown(obj))

        # Extract images from DOCX archive
        images = self._extract_images(file_path, doc_output_dir)
        for image_path in images:
            md_parts.append(f"![Image]({image_path})")

        md_text = "\n\n".join(md_parts)

        out_md = doc_output_dir / f"{file_path.stem}.md"
        out_md.write_text(md_text, encoding="utf-8")
        logger.info("Converted DOCX: %s", file_path.name)

        return Document(
            source_path=str(file_path),
            title=title or file_path.stem,
            format="docx",
            raw_content=md_text,
            metadata={"output_dir": str(doc_output_dir)},
        )

    @staticmethod
    def _table_to_markdown(table: Table) -> str:
        """Convert a python-docx Table to a markdown table string."""
        rows: list[str] = []
        num_cols = 0

        for i, row in enumerate(table.rows):
            cells = [
                cell.text.strip().replace("|", "\\|") for cell in row.cells
            ]
            num_cols = max(num_cols, len(cells))
            rows.append("| " + " | ".join(cells) + " |")

            if i == 0:
                rows.append("| " + " | ".join(["---"] * num_cols) + " |")

        return "\n".join(rows)

    @staticmethod
    def _extract_images(docx_path: Path, output_dir: Path) -> list[str]:
        """Extract embedded images from the DOCX ZIP archive."""
        images_dir = output_dir / "images"
        extracted: list[str] = []

        try:
            with zipfile.ZipFile(docx_path, "r") as zf:
                media_files = [
                    n for n in zf.namelist() if n.startswith("word/media/")
                ]
                if not media_files:
                    return []

                images_dir.mkdir(parents=True, exist_ok=True)
                for name in media_files:
                    image_filename = Path(name).name
                    image_data = zf.read(name)
                    target = images_dir / image_filename
                    target.write_bytes(image_data)
                    extracted.append(f"images/{image_filename}")
        except zipfile.BadZipFile:
            logger.warning(
                "Could not extract images from %s (bad zip)", docx_path
            )

        return extracted
